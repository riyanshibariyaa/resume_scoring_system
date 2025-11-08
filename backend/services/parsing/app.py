"""
Resume Parsing Service
Extracts text and structure from PDF, DOCX, and TXT files
"""

import os
import io
import json
import hashlib
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime

from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from loguru import logger
import fitz  # PyMuPDF
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.layout import LAParams
import docx
import pytesseract
from PIL import Image

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = os.getenv('UPLOAD_FOLDER', './uploads')
PARSED_FOLDER = os.getenv('PARSED_FOLDER', './parsed')
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'rtf'}

# Create directories
Path(UPLOAD_FOLDER).mkdir(parents=True, exist_ok=True)
Path(PARSED_FOLDER).mkdir(parents=True, exist_ok=True)

# Configure logger
logger.add("parsing_service.log", rotation="10 MB", retention="30 days", level="INFO")


class DocumentParser:
    """Main document parsing class"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt', 'rtf']
    
    def parse_document(self, file_path: str, file_format: str) -> Dict[str, Any]:
        """
        Parse document based on format
        
        Args:
            file_path: Path to the document
            file_format: Format of the document (pdf, docx, txt)
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        try:
            logger.info(f"Parsing document: {file_path} (format: {file_format})")
            
            if file_format == 'pdf':
                return self._parse_pdf(file_path)
            elif file_format == 'docx':
                return self._parse_docx(file_path)
            elif file_format in ['txt', 'rtf']:
                return self._parse_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_format}")
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF document"""
        result = {
            'text': '',
            'pages': [],
            'metadata': {},
            'sections': [],
            'parsing_method': 'text',
            'ocr_used': False
        }
        
        try:
            # Try PyMuPDF first (faster)
            doc = fitz.open(file_path)
            result['metadata'] = {
                'pages': doc.page_count,
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'keywords': doc.metadata.get('keywords', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'creation_date': doc.metadata.get('creationDate', ''),
            }
            
            full_text = []
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                # If no text found, try OCR
                if len(page_text.strip()) < 50:
                    logger.info(f"Low text content on page {page_num}, attempting OCR")
                    page_text = self._ocr_page(page)
                    result['ocr_used'] = True
                
                full_text.append(page_text)
                result['pages'].append({
                    'page_number': page_num + 1,
                    'text': page_text,
                    'char_count': len(page_text)
                })
            
            result['text'] = '\n\n'.join(full_text)
            doc.close()
            
            # If text extraction failed, try pdfminer as fallback
            if len(result['text'].strip()) < 100:
                logger.info("Low text extraction, trying pdfminer fallback")
                result['text'] = pdf_extract_text(file_path, laparams=LAParams())
                result['parsing_method'] = 'pdfminer'
            
            # Detect sections
            result['sections'] = self._detect_sections(result['text'])
            
            logger.info(f"Successfully parsed PDF: {len(result['text'])} characters")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX document"""
        result = {
            'text': '',
            'paragraphs': [],
            'metadata': {},
            'sections': [],
            'parsing_method': 'python-docx',
            'ocr_used': False
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract core properties
            core_props = doc.core_properties
            result['metadata'] = {
                'author': core_props.author if core_props.author else '',
                'title': core_props.title if core_props.title else '',
                'subject': core_props.subject if core_props.subject else '',
                'keywords': core_props.keywords if core_props.keywords else '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            }
            
            # Extract text
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                    result['paragraphs'].append({
                        'text': text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            result['text'] = '\n'.join(paragraphs)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        result['text'] += '\n' + row_text
            
            # Detect sections
            result['sections'] = self._detect_sections(result['text'])
            
            logger.info(f"Successfully parsed DOCX: {len(result['text'])} characters")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise
    
    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """Parse plain text document"""
        result = {
            'text': '',
            'lines': [],
            'metadata': {},
            'sections': [],
            'parsing_method': 'text',
            'ocr_used': False
        }
        
        try:
            # Detect encoding
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] or 'utf-8'
            
            # Read file
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                text = f.read()
            
            result['text'] = text
            result['lines'] = [line.strip() for line in text.split('\n') if line.strip()]
            result['metadata'] = {
                'encoding': encoding,
                'line_count': len(result['lines']),
                'char_count': len(text)
            }
            
            # Detect sections
            result['sections'] = self._detect_sections(text)
            
            logger.info(f"Successfully parsed text file: {len(text)} characters")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing text file: {str(e)}")
            raise
    
    def _ocr_page(self, page) -> str:
        """Perform OCR on a PDF page"""
        try:
            # Convert page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
            img_data = pix.tobytes("png")
            
            # OCR using Tesseract
            image = Image.open(io.BytesIO(img_data))
            text = pytesseract.image_to_string(image, lang='eng')
            
            return text
            
        except Exception as e:
            logger.warning(f"OCR failed: {str(e)}")
            return ""
    
    def _detect_sections(self, text: str) -> List[Dict[str, Any]]:
        """
        Detect common resume sections using pattern matching
        
        Returns list of detected sections with their positions
        """
        sections = []
        
        # Common section headers (case-insensitive patterns)
        section_patterns = {
            'contact': ['contact', 'contact information', 'personal information', 'personal details'],
            'summary': ['summary', 'profile', 'objective', 'about me', 'professional summary'],
            'experience': ['experience', 'work experience', 'employment', 'work history', 'professional experience'],
            'education': ['education', 'academic', 'qualifications'],
            'skills': ['skills', 'technical skills', 'competencies', 'expertise', 'core competencies'],
            'certifications': ['certifications', 'certificates', 'licenses', 'professional certifications'],
            'projects': ['projects', 'key projects', 'notable projects'],
            'awards': ['awards', 'honors', 'achievements', 'recognition'],
            'publications': ['publications', 'papers', 'research'],
            'languages': ['languages', 'language skills']
        }
        
        text_lower = text.lower()
        lines = text.split('\n')
        
        for section_type, patterns in section_patterns.items():
            for i, line in enumerate(lines):
                line_lower = line.lower().strip()
                
                # Check if line matches any pattern
                for pattern in patterns:
                    if pattern in line_lower and len(line_lower) < 50:  # Likely a header
                        sections.append({
                            'type': section_type,
                            'header': line.strip(),
                            'line_number': i,
                            'position': text_lower.find(line_lower)
                        })
                        break
        
        # Sort by position in document
        sections.sort(key=lambda x: x['position'])
        
        return sections


def calculate_file_hash(file_path: str) -> str:
    """Calculate SHA256 hash of file"""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'service': 'parsing-service',
        'timestamp': datetime.utcnow().isoformat()
    })


@app.route('/parse', methods=['POST'])
def parse_resume():
    """
    Parse uploaded resume file
    
    Request:
        - file: Resume file (PDF, DOCX, TXT)
        - metadata: Optional JSON metadata
        
    Response:
        - parsed_data: Extracted text and structure
        - file_hash: SHA256 hash of the file
        - metadata: File metadata
    """
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({
                'error': f'File type not allowed. Supported: {", ".join(ALLOWED_EXTENSIONS)}'
            }), 400
        
        # Get metadata if provided
        metadata = {}
        if 'metadata' in request.form:
            try:
                metadata = json.loads(request.form['metadata'])
            except json.JSONDecodeError:
                logger.warning("Invalid metadata JSON provided")
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower()
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        saved_filename = f"{timestamp}_{filename}"
        file_path = os.path.join(UPLOAD_FOLDER, saved_filename)
        
        file.save(file_path)
        logger.info(f"File saved: {file_path}")
        
        # Calculate file hash
        file_hash = calculate_file_hash(file_path)
        
        # Parse the document
        parser = DocumentParser()
        parsed_data = parser.parse_document(file_path, file_extension)
        
        # Save parsed output
        parsed_filename = f"{timestamp}_{file_hash[:12]}_parsed.json"
        parsed_path = os.path.join(PARSED_FOLDER, parsed_filename)
        
        with open(parsed_path, 'w', encoding='utf-8') as f:
            json.dump(parsed_data, f, indent=2, ensure_ascii=False)
        
        # Prepare response
        response = {
            'success': True,
            'file_hash': file_hash,
            'original_filename': filename,
            'file_format': file_extension.upper(),
            'file_size': os.path.getsize(file_path),
            'parsed_data': {
                'text': parsed_data['text'],
                'char_count': len(parsed_data['text']),
                'word_count': len(parsed_data['text'].split()),
                'sections': parsed_data['sections'],
                'metadata': parsed_data['metadata'],
                'parsing_method': parsed_data['parsing_method'],
                'ocr_used': parsed_data['ocr_used']
            },
            'storage': {
                'raw_file_path': file_path,
                'parsed_file_path': parsed_path
            },
            'processed_at': datetime.utcnow().isoformat(),
            'metadata': metadata
        }
        
        logger.info(f"Successfully parsed document: {filename} (hash: {file_hash[:12]})")
        return jsonify(response), 200
        
    except Exception as e:
        logger.error(f"Error in parse_resume: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/parse/batch', methods=['POST'])
def parse_batch():
    """
    Parse multiple resume files in batch
    
    Request:
        - files: List of resume files
        
    Response:
        - results: List of parsing results
        - summary: Batch processing summary
    """
    try:
        files = request.files.getlist('files')
        
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        
        results = []
        successful = 0
        failed = 0
        
        for file in files:
            if file.filename and allowed_file(file.filename):
                try:
                    # Process each file
                    filename = secure_filename(file.filename)
                    file_extension = filename.rsplit('.', 1)[1].lower()
                    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                    saved_filename = f"{timestamp}_{filename}"
                    file_path = os.path.join(UPLOAD_FOLDER, saved_filename)
                    
                    file.save(file_path)
                    file_hash = calculate_file_hash(file_path)
                    
                    parser = DocumentParser()
                    parsed_data = parser.parse_document(file_path, file_extension)
                    
                    results.append({
                        'filename': filename,
                        'file_hash': file_hash,
                        'status': 'success',
                        'char_count': len(parsed_data['text']),
                        'sections_found': len(parsed_data['sections'])
                    })
                    successful += 1
                    
                except Exception as e:
                    results.append({
                        'filename': file.filename,
                        'status': 'failed',
                        'error': str(e)
                    })
                    failed += 1
                    logger.error(f"Failed to parse {file.filename}: {str(e)}")
        
        return jsonify({
            'success': True,
            'results': results,
            'summary': {
                'total': len(files),
                'successful': successful,
                'failed': failed
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error in parse_batch: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


if __name__ == '__main__':
    port = int(os.getenv('PORT', 5001))
    debug = os.getenv('DEBUG', 'False').lower() == 'true'
    
    logger.info(f"Starting Parsing Service on port {port}")
    app.run(host='0.0.0.0', port=port, debug=debug)
